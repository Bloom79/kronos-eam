"""
Enhanced document management service with standard documents and copying
"""

from typing import Optional, List, Dict, Any, BinaryIO
from datetime import datetime, timedelta
from sqlalchemy.orm import Session
from sqlalchemy import and_, or_, func
import os
import hashlib
import shutil
from pathlib import Path
import mimetypes
import io
from jinja2 import Template, Environment, FileSystemLoader
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document as DocxDocument
from docx.shared import Pt, Inches

from app.models.document import (
    Document, DocumentVersion, DocumentCopy, DocumentTemplate,
    WorkflowDocumentTemplate, DocumentTypeEnum, DocumentCategoryEnum, DocumentStatusEnum
)
from app.models.user import User
from app.models.workflow import WorkflowTask, Workflow
from app.models.plant import Plant
from app.models.audit import TipoModificaEnum
from app.services.audit_service import AuditService
from app.services.notification_service import NotificationService
from app.core.config import settings
from app.core.storage import StorageBackend


class DocumentService:
    """Service for comprehensive document management"""
    
    def __init__(self, db: Session):
        self.db = db
        self.storage = StorageBackend()
        self.audit_service = AuditService(db)
        self.notification_service = NotificationService(db)
    
    def create_document(
        self,
        nome: str,
        file: BinaryIO,
        categoria: DocumentCategoryEnum,
        tenant_id: int,
        user_id: int,
        descrizione: Optional[str] = None,
        impianto_id: Optional[int] = None,
        workflow_id: Optional[int] = None,
        task_id: Optional[int] = None,
        data_scadenza: Optional[datetime] = None,
        tags: Optional[List[str]] = None,
        metadata: Optional[Dict[str, Any]] = None,
        is_standard: bool = False,
        riferimenti_normativi: Optional[List[str]] = None,
        link_esterni: Optional[List[str]] = None
    ) -> Document:
        """
        Create a new document with versioning and audit trail
        
        Args:
            nome: Document name
            file: File binary content
            categoria: Document category
            tenant_id: Tenant ID
            user_id: User creating the document
            descrizione: Document description
            impianto_id: Related plant ID
            workflow_id: Related workflow ID
            task_id: Related task ID
            data_scadenza: Expiry date
            tags: Document tags
            metadata: Additional metadata
            is_standard: Whether this is a standard document
            riferimenti_normativi: Normative references
            link_esterni: External links
            
        Returns:
            Created document
        """
        # Determine file type
        filename = getattr(file, 'filename', nome)
        mime_type = mimetypes.guess_type(filename)[0]
        extension = Path(filename).suffix.upper().lstrip('.')
        
        # Map to DocumentTypeEnum
        tipo_map = {
            'PDF': DocumentTypeEnum.PDF,
            'DOC': DocumentTypeEnum.DOC,
            'DOCX': DocumentTypeEnum.DOCX,
            'XLS': DocumentTypeEnum.XLS,
            'XLSX': DocumentTypeEnum.XLSX,
            'JPG': DocumentTypeEnum.IMG,
            'JPEG': DocumentTypeEnum.IMG,
            'PNG': DocumentTypeEnum.IMG,
            'XML': DocumentTypeEnum.XML,
            'P7M': DocumentTypeEnum.P7M
        }
        tipo = tipo_map.get(extension, DocumentTypeEnum.PDF)
        
        # Read file content
        content = file.read()
        file_size = len(content)
        
        # Calculate checksum
        checksum = hashlib.sha256(content).hexdigest()
        
        # Store file
        file_path = self.storage.store_file(
            content=content,
            filename=filename,
            tenant_id=tenant_id,
            subfolder=f"documents/{categoria.value.lower()}"
        )
        
        # Create document record
        document = Document(
            nome=nome,
            descrizione=descrizione,
            tipo=tipo,
            categoria=categoria,
            stato=DocumentStatusEnum.VALIDO,
            file_path=file_path,
            file_size=file_size,
            mime_type=mime_type,
            checksum=checksum,
            impianto_id=impianto_id,
            workflow_id=workflow_id,
            task_id=task_id,
            data_scadenza=data_scadenza,
            tags=tags or [],
            model_metadata=metadata or {},
            is_standard=is_standard,
            riferimenti_normativi=riferimenti_normativi or [],
            link_esterni=link_esterni or [],
            tenant_id=tenant_id
        )
        
        self.db.add(document)
        self.db.commit()
        self.db.refresh(document)
        
        # Create initial version
        self._create_version(document, user_id, "Versione iniziale")
        
        # Log creation
        self.audit_service.log_change(
            entity_type="document",
            entity_id=document.id,
            tipo_modifica=TipoModificaEnum.CREAZIONE,
            user_id=user_id,
            tenant_id=tenant_id,
            new_values={
                "nome": nome,
                "categoria": categoria.value,
                "is_standard": is_standard
            }
        )
        
        return document
    
    def update_document(
        self,
        document_id: int,
        user_id: int,
        tenant_id: int,
        nome: Optional[str] = None,
        descrizione: Optional[str] = None,
        data_scadenza: Optional[datetime] = None,
        tags: Optional[List[str]] = None,
        metadata: Optional[Dict[str, Any]] = None,
        riferimenti_normativi: Optional[List[str]] = None,
        link_esterni: Optional[List[str]] = None,
        new_file: Optional[BinaryIO] = None,
        version_note: Optional[str] = None
    ) -> Document:
        """Update document with version tracking"""
        document = self.db.query(Document).filter(
            Document.id == document_id,
            Document.tenant_id == tenant_id
        ).first()
        
        if not document:
            raise ValueError("Document not found")
        
        # Capture old values for audit
        old_values = {
            "nome": document.nome,
            "descrizione": document.descrizione,
            "data_scadenza": document.data_scadenza.isoformat() if document.data_scadenza else None
        }
        
        # Update fields
        if nome is not None:
            document.nome = nome
        if descrizione is not None:
            document.descrizione = descrizione
        if data_scadenza is not None:
            document.data_scadenza = data_scadenza
        if tags is not None:
            document.tags = tags
        if metadata is not None:
            document.model_metadata.update(metadata)
        if riferimenti_normativi is not None:
            document.riferimenti_normativi = riferimenti_normativi
        if link_esterni is not None:
            document.link_esterni = link_esterni
        
        # Handle new file upload
        if new_file:
            # Create new version before updating
            self._create_version(document, user_id, version_note or "File aggiornato")
            
            # Store new file
            content = new_file.read()
            file_size = len(content)
            checksum = hashlib.sha256(content).hexdigest()
            
            filename = getattr(new_file, 'filename', document.nome)
            file_path = self.storage.store_file(
                content=content,
                filename=filename,
                tenant_id=tenant_id,
                subfolder=f"documents/{document.categoria.value.lower()}"
            )
            
            # Update document
            document.file_path = file_path
            document.file_size = file_size
            document.checksum = checksum
            document.versione += 1
        
        document.data_ultima_modifica = datetime.utcnow()
        
        # Capture new values
        new_values = {
            "nome": document.nome,
            "descrizione": document.descrizione,
            "data_scadenza": document.data_scadenza.isoformat() if document.data_scadenza else None
        }
        
        self.db.commit()
        self.db.refresh(document)
        
        # Log update
        self.audit_service.log_change(
            entity_type="document",
            entity_id=document.id,
            tipo_modifica=TipoModificaEnum.AGGIORNAMENTO,
            user_id=user_id,
            tenant_id=tenant_id,
            old_values=old_values,
            new_values=new_values,
            note=version_note
        )
        
        # TODO: Notify if standard document updated (requires async handling)
        # if document.is_standard:
        #     self._notify_standard_document_update(document, user_id, tenant_id)
        
        return document
    
    def copy_document(
        self,
        document_id: int,
        user_id: int,
        tenant_id: int,
        nome_copia: Optional[str] = None,
        target_impianto_id: Optional[int] = None,
        customizations: Optional[Dict[str, Any]] = None,
        note: Optional[str] = None
    ) -> Document:
        """
        Create a customized copy of a document
        
        Args:
            document_id: Source document ID
            user_id: User creating the copy
            tenant_id: Tenant ID
            nome_copia: Name for the copy
            target_impianto_id: Target plant ID
            customizations: Customization parameters
            note: Notes about customization
            
        Returns:
            New document copy
        """
        # Get source document
        source = self.db.query(Document).filter(
            Document.id == document_id,
            Document.tenant_id == tenant_id
        ).first()
        
        if not source:
            raise ValueError("Source document not found")
        
        # Copy file
        source_content = self.storage.retrieve_file(source.file_path)
        
        # Apply customizations if needed (e.g., template processing)
        if customizations and source.tipo in [DocumentTypeEnum.DOC, DocumentTypeEnum.DOCX]:
            # This would process templates with customizations
            # For now, just copy as-is
            content = source_content
        else:
            content = source_content
        
        # Create new document
        nome = nome_copia or f"Copia di {source.nome}"
        
        new_file_path = self.storage.store_file(
            content=content,
            filename=nome,
            tenant_id=tenant_id,
            subfolder=f"documents/{source.categoria.value.lower()}/copies"
        )
        
        # Create document record
        copy = Document(
            nome=nome,
            descrizione=f"Copia personalizzata di: {source.nome}",
            tipo=source.tipo,
            categoria=source.categoria,
            stato=DocumentStatusEnum.VALIDO,
            file_path=new_file_path,
            file_size=len(content),
            mime_type=source.mime_type,
            checksum=hashlib.sha256(content).hexdigest(),
            impianto_id=target_impianto_id or source.impianto_id,
            tags=source.tags.copy() if source.tags else [],
            is_standard=False,  # Copies are not standard
            tenant_id=tenant_id
        )
        
        self.db.add(copy)
        self.db.commit()
        self.db.refresh(copy)
        
        # Create copy record
        copy_record = DocumentCopy(
            documento_originale_id=source.id,
            utente_creazione_id=user_id,
            nome_copia=nome,
            contenuto_customizzato=str(customizations) if customizations else None,
            modifiche_applicate=customizations or {},
            note_personalizzazione=note,
            tenant_id=tenant_id
        )
        
        self.db.add(copy_record)
        self.db.commit()
        
        # Log copy creation
        self.audit_service.log_change(
            entity_type="document",
            entity_id=copy.id,
            tipo_modifica=TipoModificaEnum.CREAZIONE,
            user_id=user_id,
            tenant_id=tenant_id,
            dettaglio_modifica={
                "source_document_id": source.id,
                "is_copy": True
            }
        )
        
        return copy
    
    def search_documents(
        self,
        tenant_id: int,
        query: Optional[str] = None,
        categoria: Optional[DocumentCategoryEnum] = None,
        tipo: Optional[DocumentTypeEnum] = None,
        stato: Optional[DocumentStatusEnum] = None,
        impianto_id: Optional[int] = None,
        workflow_id: Optional[int] = None,
        tags: Optional[List[str]] = None,
        is_standard: Optional[bool] = None,
        riferimento_normativo: Optional[str] = None,
        data_scadenza_start: Optional[datetime] = None,
        data_scadenza_end: Optional[datetime] = None,
        limit: int = 50,
        offset: int = 0,
        order_by: str = "created_at",
        order_desc: bool = True
    ) -> Dict[str, Any]:
        """
        Advanced document search with multiple filters
        
        Returns:
            Dictionary with documents and metadata
        """
        base_query = self.db.query(Document).filter(
            Document.tenant_id == tenant_id
        )
        
        # Apply filters
        if query:
            # Full-text search on nome and descrizione
            base_query = base_query.filter(
                or_(
                    Document.nome.ilike(f"%{query}%"),
                    Document.descrizione.ilike(f"%{query}%")
                )
            )
        
        if categoria:
            base_query = base_query.filter(Document.categoria == categoria)
        
        if tipo:
            base_query = base_query.filter(Document.tipo == tipo)
        
        if stato:
            base_query = base_query.filter(Document.stato == stato)
        
        if impianto_id:
            base_query = base_query.filter(Document.impianto_id == impianto_id)
        
        if workflow_id:
            base_query = base_query.filter(Document.workflow_id == workflow_id)
        
        if tags:
            # Filter documents that have any of the specified tags
            for tag in tags:
                base_query = base_query.filter(Document.tags.contains([tag]))
        
        if is_standard is not None:
            base_query = base_query.filter(Document.is_standard == is_standard)
        
        if riferimento_normativo:
            base_query = base_query.filter(
                Document.riferimenti_normativi.contains([riferimento_normativo])
            )
        
        if data_scadenza_start:
            base_query = base_query.filter(Document.data_scadenza >= data_scadenza_start)
        
        if data_scadenza_end:
            base_query = base_query.filter(Document.data_scadenza <= data_scadenza_end)
        
        # Get total count
        total = base_query.count()
        
        # Apply ordering
        order_column = getattr(Document, order_by, Document.created_at)
        if order_desc:
            base_query = base_query.order_by(order_column.desc())
        else:
            base_query = base_query.order_by(order_column)
        
        # Apply pagination
        documents = base_query.offset(offset).limit(limit).all()
        
        # Get facets for filtering
        facets = self._get_search_facets(tenant_id)
        
        return {
            "documents": documents,
            "total": total,
            "limit": limit,
            "offset": offset,
            "facets": facets
        }
    
    def get_expiring_documents(
        self,
        tenant_id: int,
        days_ahead: int = 30,
        categoria: Optional[DocumentCategoryEnum] = None
    ) -> List[Document]:
        """Get documents expiring within specified days"""
        expiry_date = datetime.utcnow() + timedelta(days=days_ahead)
        
        query = self.db.query(Document).filter(
            Document.tenant_id == tenant_id,
            Document.data_scadenza <= expiry_date,
            Document.data_scadenza >= datetime.utcnow(),
            Document.stato == DocumentStatusEnum.VALIDO
        )
        
        if categoria:
            query = query.filter(Document.categoria == categoria)
        
        return query.order_by(Document.data_scadenza).all()
    
    def get_standard_documents(
        self,
        tenant_id: int,
        categoria: Optional[DocumentCategoryEnum] = None
    ) -> List[Document]:
        """Get all standard documents available for copying"""
        query = self.db.query(Document).filter(
            Document.tenant_id == tenant_id,
            Document.is_standard == True,
            Document.stato == DocumentStatusEnum.VALIDO
        )
        
        if categoria:
            query = query.filter(Document.categoria == categoria)
        
        return query.order_by(Document.nome).all()
    
    def link_document_to_task(
        self,
        document_id: int,
        task_id: int,
        user_id: int,
        tenant_id: int
    ) -> bool:
        """Link a document to a workflow task"""
        document = self.db.query(Document).filter(
            Document.id == document_id,
            Document.tenant_id == tenant_id
        ).first()
        
        task = self.db.query(WorkflowTask).filter(
            WorkflowTask.id == task_id,
            WorkflowTask.workflow.has(tenant_id=tenant_id)
        ).first()
        
        if not document or not task:
            return False
        
        # Update document
        document.task_id = task_id
        
        # Update task's associated documents
        if not task.documenti_associati:
            task.documenti_associati = []
        
        if document_id not in task.documenti_associati:
            task.documenti_associati.append(document_id)
        
        self.db.commit()
        
        # Log the linking
        self.audit_service.log_change(
            entity_type="document",
            entity_id=document_id,
            tipo_modifica=TipoModificaEnum.AGGIORNAMENTO,
            user_id=user_id,
            tenant_id=tenant_id,
            dettaglio_modifica={
                "action": "linked_to_task",
                "task_id": task_id
            }
        )
        
        return True
    
    def _create_version(
        self,
        document: Document,
        user_id: int,
        note: str
    ):
        """Create a document version record"""
        user = self.db.query(User).filter(User.id == user_id).first()
        
        version = DocumentVersion(
            document_id=document.id,
            versione=document.versione,
            file_path=document.file_path,
            file_size=document.file_size,
            checksum=document.checksum,
            modifiche=note,
            modificato_da=user.email if user else "Sistema",
            tenant_id=document.tenant_id
        )
        
        self.db.add(version)
        self.db.commit()
    
    def _get_search_facets(self, tenant_id: int) -> Dict[str, Any]:
        """Get facets for search filtering"""
        base_query = self.db.query(Document).filter(
            Document.tenant_id == tenant_id
        )
        
        # Category counts
        category_counts = {}
        for categoria in DocumentCategoryEnum:
            count = base_query.filter(Document.categoria == categoria).count()
            category_counts[categoria.value] = count
        
        # Type counts
        type_counts = {}
        for tipo in DocumentTypeEnum:
            count = base_query.filter(Document.tipo == tipo).count()
            type_counts[tipo.value] = count
        
        # Status counts
        status_counts = {}
        for stato in DocumentStatusEnum:
            count = base_query.filter(Document.stato == stato).count()
            status_counts[stato.value] = count
        
        # Get unique tags
        all_tags = self.db.query(
            func.unnest(Document.tags).label('tag')
        ).filter(
            Document.tenant_id == tenant_id
        ).distinct().all()
        
        tags = [tag[0] for tag in all_tags if tag[0]]
        
        return {
            "categories": category_counts,
            "types": type_counts,
            "statuses": status_counts,
            "tags": tags,
            "has_standard_docs": base_query.filter(Document.is_standard == True).count() > 0,
            "has_expiring_docs": base_query.filter(
                Document.data_scadenza <= datetime.utcnow() + timedelta(days=30)
            ).count() > 0
        }
    
    async def _notify_standard_document_update(
        self,
        document: Document,
        user_id: int,
        tenant_id: int
    ):
        """Notify users who have copies of a standard document"""
        # Find all copies
        copies = self.db.query(DocumentCopy).filter(
            DocumentCopy.documento_originale_id == document.id
        ).all()
        
        # Notify copy owners
        notified_users = set()
        for copy in copies:
            if copy.utente_creazione_id not in notified_users:
                await self.notification_service.send_notification(
                    user_id=copy.utente_creazione_id,
                    tipo="documento",
                    titolo=f"Documento standard aggiornato: {document.nome}",
                    messaggio=f"Il documento standard '{document.nome}' è stato aggiornato. Verifica se è necessario aggiornare la tua copia personalizzata.",
                    tenant_id=tenant_id,
                    documento_id=document.id,
                    link=f"/documents/{document.id}",
                    metadata={
                        "update_type": "standard_document_update",
                        "copy_count": len(copies)
                    }
                )
                notified_users.add(copy.utente_creazione_id)
    
    def generate_from_template(
        self,
        template_id: int,
        data: Dict[str, Any],
        output_format: str,
        user_id: int,
        tenant_id: int,
        workflow_id: Optional[int] = None,
        task_id: Optional[int] = None,
        impianto_id: Optional[int] = None
    ) -> Document:
        """
        Generate a document from a template
        
        Args:
            template_id: Document template ID
            data: Data to fill in the template
            output_format: Output format (pdf, docx)
            user_id: User generating the document
            tenant_id: Tenant ID
            workflow_id: Associated workflow ID
            task_id: Associated task ID
            impianto_id: Associated plant ID
            
        Returns:
            Generated document
        """
        # Get template
        template = self.db.query(DocumentTemplate).filter(
            DocumentTemplate.id == template_id,
            DocumentTemplate.tenant_id == tenant_id,
            DocumentTemplate.attivo == True
        ).first()
        
        if not template:
            raise ValueError("Template not found or inactive")
        
        # Load additional data if workflow/impianto provided
        context = data.copy()
        
        if workflow_id:
            workflow = self.db.query(Workflow).filter(
                Workflow.id == workflow_id,
                Workflow.tenant_id == tenant_id
            ).first()
            if workflow:
                context['workflow'] = {
                    'nome': workflow.nome,
                    'tipo': workflow.tipo,
                    'stato': workflow.stato_corrente,
                    'data_inizio': workflow.data_inizio.isoformat() if workflow.data_inizio else None
                }
        
        if impianto_id:
            impianto = self.db.query(Plant).filter(
                Plant.id == impianto_id,
                Plant.tenant_id == tenant_id
            ).first()
            if impianto:
                context['impianto'] = {
                    'nome': impianto.nome,
                    'tipo': impianto.tipo,
                    'potenza_kw': impianto.potenza_kw,
                    'indirizzo': impianto.indirizzo,
                    'comune': impianto.comune,
                    'provincia': impianto.provincia,
                    'partita_iva': impianto.partita_iva,
                    'codice_fiscale': impianto.codice_fiscale
                }
        
        # Add current date
        context['data_generazione'] = datetime.now().strftime("%d/%m/%Y")
        
        # Generate document based on format
        if output_format.lower() == 'pdf':
            content = self._generate_pdf_from_template(template, context)
            tipo = DocumentTypeEnum.PDF
            mime_type = 'application/pdf'
            extension = '.pdf'
        elif output_format.lower() == 'docx':
            content = self._generate_docx_from_template(template, context)
            tipo = DocumentTypeEnum.DOCX
            mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            extension = '.docx'
        else:
            raise ValueError(f"Unsupported output format: {output_format}")
        
        # Create filename
        filename = f"{template.nome}_{datetime.now().strftime('%Y%m%d_%H%M%S')}{extension}"
        
        # Store file
        file_path = self.storage.store_file(
            content=content,
            filename=filename,
            tenant_id=tenant_id,
            subfolder=f"documents/{template.categoria.value.lower()}/generated"
        )
        
        # Create document record
        document = Document(
            nome=filename,
            descrizione=f"Generato da template: {template.nome}",
            tipo=tipo,
            categoria=template.categoria,
            stato=DocumentStatusEnum.VALIDO,
            file_path=file_path,
            file_size=len(content),
            mime_type=mime_type,
            checksum=hashlib.sha256(content).hexdigest(),
            impianto_id=impianto_id,
            workflow_id=workflow_id,
            task_id=task_id,
            tenant_id=tenant_id,
            generated_from_template_id=template_id,
            template_data=context
        )
        
        self.db.add(document)
        self.db.commit()
        self.db.refresh(document)
        
        # Log generation
        self.audit_service.log_change(
            entity_type="document",
            entity_id=document.id,
            tipo_modifica=TipoModificaEnum.CREAZIONE,
            user_id=user_id,
            tenant_id=tenant_id,
            dettaglio_modifica={
                "action": "generated_from_template",
                "template_id": template_id,
                "output_format": output_format
            }
        )
        
        return document
    
    def _generate_pdf_from_template(self, template: DocumentTemplate, context: Dict[str, Any]) -> bytes:
        """Generate PDF from template using ReportLab"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        story = []
        styles = getSampleStyleSheet()
        
        # Load template content
        if template.template_path:
            template_dir = Path(settings.TEMPLATE_DIR) / "documents"
            env = Environment(loader=FileSystemLoader(str(template_dir)))
            jinja_template = env.get_template(template.template_path)
            content = jinja_template.render(**context)
        else:
            # Use default template
            content = self._get_default_template_content(template, context)
        
        # Parse content and build PDF
        for paragraph in content.split('\n\n'):
            if paragraph.strip():
                if paragraph.startswith('#'):
                    # Header
                    level = paragraph.count('#')
                    text = paragraph.lstrip('#').strip()
                    if level == 1:
                        para = Paragraph(text, styles['Title'])
                    elif level == 2:
                        para = Paragraph(text, styles['Heading1'])
                    else:
                        para = Paragraph(text, styles['Heading2'])
                else:
                    # Normal paragraph
                    para = Paragraph(paragraph, styles['BodyText'])
                
                story.append(para)
                story.append(Spacer(1, 0.2*inch))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.read()
    
    def _generate_docx_from_template(self, template: DocumentTemplate, context: Dict[str, Any]) -> bytes:
        """Generate DOCX from template using python-docx"""
        doc = DocxDocument()
        
        # Load template content
        if template.template_path:
            template_dir = Path(settings.TEMPLATE_DIR) / "documents"
            env = Environment(loader=FileSystemLoader(str(template_dir)))
            jinja_template = env.get_template(template.template_path)
            content = jinja_template.render(**context)
        else:
            # Use default template
            content = self._get_default_template_content(template, context)
        
        # Parse content and build DOCX
        for paragraph in content.split('\n\n'):
            if paragraph.strip():
                if paragraph.startswith('#'):
                    # Header
                    level = paragraph.count('#')
                    text = paragraph.lstrip('#').strip()
                    p = doc.add_heading(text, level)
                else:
                    # Normal paragraph
                    doc.add_paragraph(paragraph)
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
    
    def _get_default_template_content(self, template: DocumentTemplate, context: Dict[str, Any]) -> str:
        """Get default template content for connection request"""
        if template.uso == "Richiesta Connessione":
            return self._get_connection_request_template(context)
        else:
            # Generic template
            return f"""
# {template.nome}

Data: {context.get('data_generazione', '')}

## Dati Richiedente
Nome: {context.get('nome_richiedente', '')}
Codice Fiscale: {context.get('codice_fiscale', '')}

## Dati Plant
{context.get('impianto', {}).get('nome', '')}
Potenza: {context.get('impianto', {}).get('potenza_kw', '')} kW
Indirizzo: {context.get('impianto', {}).get('indirizzo', '')}
"""
    
    def _get_connection_request_template(self, context: Dict[str, Any]) -> str:
        """Get connection request template content"""
        impianto = context.get('impianto', {})
        richiedente = context.get('richiedente', {})
        
        return f"""
# Plant Fotovoltaico Intestato a

Nome e Cognome: {richiedente.get('nome', '')} {richiedente.get('cognome', '')} nato/a a {richiedente.get('luogo_nascita', '_________________')} ({richiedente.get('provincia_nascita', '')}) il {richiedente.get('data_nascita', '______________')}

Residente in {richiedente.get('indirizzo_residenza', '____________________')} n. {richiedente.get('civico', '')} Comune di {richiedente.get('comune_residenza', '_______________________________________')} Prov. ({richiedente.get('provincia_residenza', '__')})

{'In qualità di ' + richiedente.get('qualifica', '') + ' del/della ' + richiedente.get('ragione_sociale', '') if richiedente.get('tipo') == 'azienda' else ''}

## AUTORIZZAZIONI AMMINISTRATIVE

L'installazione di un impianto fotovoltaico rientra in attività di edilizia libera, come confermano sia il Testo Unico per l'Edilizia (DPR 380/2001), sia il DM 2 marzo 2018.

Titoli Autorizzativi richiesti:
{'☑' if context.get('edilizia_libera', True) else '☐'} Attività in edilizia libera.
{'☑' if context.get('altro_titolo', False) else '☐'} Altro titolo autorizzativo: {context.get('altro_titolo_desc', '____________________________________________________')}

## ITER CONNESSIONE ALLA RETE

Documenti necessari per la FASE 1 (Domanda di connessione):

L'impianto usufruisce del Superbonus 110%? SI {'☑' if context.get('superbonus', False) else '☐'} NO {'☑' if not context.get('superbonus', False) else '☐'}

Si richiedono:
☑ copia dell'ultima bolletta elettrica
☑ codice IBAN dell'intestatario della bolletta elettrica: {richiedente.get('iban', '_______________________________')}
☑ copia del documento di identità dell'intestatario della bolletta che sarà l'intestatario dell'impianto.
☑ indirizzo e.mail per la registrazione sul portale GSE: {richiedente.get('email', '________________________________')}
☑ numero di telefono: {richiedente.get('telefono', '______ ___________________________________________________')}
☑ mappa catastale che deve riportare il Comune di rilascio, il foglio e particella catastale con evidenziata la particella interessata.
☑ l'intestatario della bolletta è titolare del sito oggetto dell'installazione dell'impianto: SI {'☑' if richiedente.get('titolare_sito', True) else '☐'} NO {'☑' if not richiedente.get('titolare_sito', True) else '☐'}

## Dati Plant
Potenza: {impianto.get('potenza_kw', '')} kW
Ubicazione: {impianto.get('indirizzo', '')} - {impianto.get('comune', '')} ({impianto.get('provincia', '')})

## Corrispettivo Preventivo
{self._calculate_fee_text(impianto.get('potenza_kw', 0))}

Data: {context.get('data_generazione', '')}

Firma del Richiedente: _______________________________
"""
    
    def _calculate_fee_text(self, potenza_kw: float) -> str:
        """Calculate connection fee based on power"""
        if potenza_kw <= 6:
            return "Potenze fino a 6 kW: € 30 + IVA 22% = € 36,60"
        elif potenza_kw <= 10:
            return "Potenze da 6 a 10 kW: € 50 + IVA 22% = € 61,00"
        elif potenza_kw <= 50:
            return "Potenze da 10 a 50 kW: € 100 + IVA 22% = € 122,00"
        elif potenza_kw <= 100:
            return "Potenze da 50 a 100 kW: € 200 + IVA 22% = € 244,00"
        elif potenza_kw <= 500:
            return "Potenze da 100 a 500 kW: € 500 + IVA 22% = € 610,00"
        elif potenza_kw <= 1000:
            return "Potenze da 500 a 1000 kW: € 1.500 + IVA 22% = € 1.830,00"
        else:
            return "Potenze oltre 1000 kW: € 2.500 + IVA 22% = € 3.050,00"
    
    def get_template_placeholders(self, template_id: int, tenant_id: int) -> Dict[str, Any]:
        """Get required placeholders for a template"""
        template = self.db.query(DocumentTemplate).filter(
            DocumentTemplate.id == template_id,
            DocumentTemplate.tenant_id == tenant_id
        ).first()
        
        if not template:
            raise ValueError("Template not found")
        
        # Return template variables
        return template.variables or {}
    
    def get_workflow_document_templates(
        self,
        workflow_template_id: int,
        task_nome: Optional[str] = None,
        tenant_id: Optional[int] = None
    ) -> List[WorkflowDocumentTemplate]:
        """Get document templates associated with a workflow template"""
        query = self.db.query(WorkflowDocumentTemplate).filter(
            WorkflowDocumentTemplate.workflow_template_id == workflow_template_id
        )
        
        if task_nome:
            query = query.filter(WorkflowDocumentTemplate.task_nome == task_nome)
        
        if tenant_id:
            query = query.filter(WorkflowDocumentTemplate.tenant_id == tenant_id)
        
        return query.order_by(WorkflowDocumentTemplate.ordine).all()


class StorageBackend:
    """File storage backend (placeholder for actual implementation)"""
    
    def __init__(self):
        self.base_path = Path(settings.UPLOAD_PATH)
        self.base_path.mkdir(parents=True, exist_ok=True)
    
    def store_file(
        self,
        content: bytes,
        filename: str,
        tenant_id: int,
        subfolder: str = ""
    ) -> str:
        """Store file and return path"""
        # Create tenant directory
        tenant_path = self.base_path / str(tenant_id) / subfolder
        tenant_path.mkdir(parents=True, exist_ok=True)
        
        # Generate unique filename
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        unique_filename = f"{timestamp}_{filename}"
        
        file_path = tenant_path / unique_filename
        
        # Write file
        with open(file_path, 'wb') as f:
            f.write(content)
        
        # Return relative path
        return str(file_path.relative_to(self.base_path))
    
    def retrieve_file(self, file_path: str) -> bytes:
        """Retrieve file content"""
        full_path = self.base_path / file_path
        
        with open(full_path, 'rb') as f:
            return f.read()
    
    def delete_file(self, file_path: str):
        """Delete file"""
        full_path = self.base_path / file_path
        
        if full_path.exists():
            full_path.unlink()